# convert_logs_to_docx.py

import os
import json
from docx import Document
from docx.shared import Pt
from pathlib import Path

BASE_DIR = Path("/Users/libertyelectronics/Desktop/curriculum_builder/CB_Agent")
LOGS_DIR = BASE_DIR / "logs"
OUTPUT_DIR = BASE_DIR / "genai_output"
OUTPUT_DIR.mkdir(exist_ok=True)

def load_inputs():
    """Load the user inputs from JSON file."""
    with open("user_inputs.json", "r") as f:
        inputs = json.load(f)
    return inputs

def convert_md_to_docx(subject: str, grade_level: str, topic: str):
    """Convert all .md files in the logs folder to .docx files in genai_output."""
    for md_file in LOGS_DIR.glob("*.md"):
        agent_name = md_file.stem.split("_task")[0]
        
        # Define the output filename based on inputs
        output_filename = f"{subject}_{grade_level}_{topic}_{agent_name}.docx"
        output_path = OUTPUT_DIR / output_filename
        
        with open(md_file, "r", encoding="utf-8") as f:
            md_content = f.read()
        
        doc = Document()
        doc.add_heading(f"{subject} - {grade_level} - {topic} - {agent_name}", level=1)
        
        for line in md_content.splitlines():
            line = line.strip()  # Remove leading and trailing whitespace
            if line.startswith("**") and line.endswith("**"):
                # Bold Text (remove the ** syntax)
                paragraph = doc.add_paragraph()
                paragraph.add_run(line.strip("**")).bold = True
            elif "|" in line:
                # Table Row
                table_data = [cell.strip() for cell in line.split("|") if cell.strip()]
                if len(table_data) > 1:  # Ignore separator lines like ----
                    if not doc.tables:
                        # Create table if it doesn't exist
                        table = doc.add_table(rows=1, cols=len(table_data))
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        for i, cell_text in enumerate(table_data):
                            hdr_cells[i].text = cell_text
                    else:
                        # Add row to existing table
                        row_cells = doc.tables[0].add_row().cells
                        for i, cell_text in enumerate(table_data):
                            row_cells[i].text = cell_text
            elif line.startswith("- **"):
                # Bold label with content (e.g., - **Subject**: Content)
                label, content = line.split(":", 1)
                paragraph = doc.add_paragraph()
                paragraph.add_run(label.strip("- ")).bold = True
                paragraph.add_run(": " + content.strip())
            elif line.startswith("#"):
                # Heading (remove the # syntax)
                heading_level = min(3, line.count("#"))  # Use only up to Heading 3
                doc.add_heading(line.strip("# "), level=heading_level)
            elif line.startswith("**") and "**" in line[2:]:
                # Inline bold text (e.g., "This is **bold** text.")
                parts = line.split("**")
                paragraph = doc.add_paragraph()
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Odd indices are bold text
                        paragraph.add_run(part).bold = True
                    else:
                        paragraph.add_run(part)
            else:
                # Regular paragraph for lines without special formatting
                doc.add_paragraph(line)
        
        doc.save(str(output_path))  # Convert Path to string for save compatibility
        print(f"Converted '{md_file.name}' to '{output_filename}'")

    print("All log files have been converted successfully.")

if __name__ == "__main__":
    # Load inputs for naming files
    inputs = load_inputs()
    subject = inputs['subject'].replace(" ", "_").lower()
    grade_level = inputs['grade_level'].replace(" ", "_").lower()
    topic = inputs['topic'].replace(" ", "_").lower()

    convert_md_to_docx(subject, grade_level, topic)
